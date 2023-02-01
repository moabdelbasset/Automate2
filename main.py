from flask import Flask, request, render_template, send_file
import os
import docx

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    title = request.form["title"]
    content = request.form["content"]

    filename = "generated_file.docx"
    file_path = os.path.join("files", filename)

    doc = docx.Document()
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(content)
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    app.run()
